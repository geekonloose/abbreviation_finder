'''
File: acronym_finder.py
Project: 
File Created: Thursday, 31st October 2019 9:26:15 am
Author: John Arul & Parth (arul@igcar.gov.in, parthpdpu@gmail.com)
-----
Last Modified: Friday, 1st November 2019 9:02:12 am
Modified By: John Arul & Parth Patel (arul@igcar.gov.in, parthpdpu@gmail.com)
-----
Copyright: IGCAR - 2019
'''

import docx
import re
print('Enter name of file with extension:\n')
filename = input()
print('enter min len of acronym you want to incorporate in search:\n')
min_no = int(input())

set_acronym = set()
if filename.split('.')[1] == 'docx':
    doc = docx.Document(filename)
    for i in range(len(doc.paragraphs)):
        lst_acronym = re.findall('[A-Z]*', doc.paragraphs[i].text)
        for j in lst_acronym:
            if len(j) >= min_no:
                set_acronym.add(j)
else:
    file = open(filename, 'r')
    data = file.readlines()
    for line in data:
        lst_acronym = re.findall('[A-Z]*', line)
        for j in lst_acronym:
            if len(j) >= min_no:
                set_acronym.add(j)
f = open('acronym.txt', 'w')
for i in sorted(set_acronym):
    f.write(f'{i}\n')
f.close()
